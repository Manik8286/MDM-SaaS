"""
Generate MDM SaaS Local Deployment Guide as a Word document.
Run: python3 scripts/gen_deployment_doc.py
Output: MDM_SaaS_Local_Deployment_Guide.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ── Colors ────────────────────────────────────────────────────────────────────
C_BLUE      = RGBColor(0x1D, 0x4E, 0xD8)
C_DARK      = RGBColor(0x0F, 0x17, 0x2A)
C_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
C_GRAY      = RGBColor(0x64, 0x74, 0x8B)
C_GREEN     = RGBColor(0x05, 0x96, 0x69)
C_ORANGE    = RGBColor(0xEA, 0x58, 0x0C)
C_CODE_BG   = RGBColor(0xF1, 0xF5, 0xF9)
C_LIGHT_BLU = RGBColor(0xDB, 0xEA, 0xFE)
C_HDR_BG    = RGBColor(0x1D, 0x4E, 0xD8)
C_ROW_ALT   = RGBColor(0xF8, 0xFA, 0xFF)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
    tcPr.append(shd)

def set_cell_border(cell, sides=('top','bottom','left','right'), color='CCCCCC', sz='4'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in sides:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), sz)
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.keep_with_next = True
    # Blue left border via shading trick — use run color + size
    run = p.add_run(text)
    run.font.size  = Pt(20)
    run.font.bold  = True
    run.font.color.rgb = C_BLUE
    # Bottom border paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '1D4ED8')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.size  = Pt(14)
    run.font.bold  = True
    run.font.color.rgb = C_DARK
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.size  = Pt(12)
    run.font.bold  = True
    run.font.color.rgb = C_BLUE
    return p

def body(text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if bold_parts:
        # bold_parts: list of (text, is_bold)
        for part, is_bold in bold_parts:
            run = p.add_run(part)
            run.font.size = Pt(11)
            run.font.bold = is_bold
            run.font.color.rgb = C_DARK
    else:
        run = p.add_run(text)
        run.font.size  = Pt(11)
        run.font.color.rgb = C_DARK
    return p

def bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent   = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_before  = Pt(1)
    p.paragraph_format.space_after   = Pt(2)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = C_DARK
        run2 = p.add_run(text)
        run2.font.size = Pt(11)
        run2.font.color.rgb = C_DARK
    else:
        run = p.add_run(text)
        run.font.size  = Pt(11)
        run.font.color.rgb = C_DARK
    return p

def numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent  = Inches(0.25)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = C_DARK
        run2 = p.add_run(text)
        run2.font.size = Pt(11)
        run2.font.color.rgb = C_DARK
    else:
        run = p.add_run(text)
        run.font.size  = Pt(11)
        run.font.color.rgb = C_DARK
    return p

def code_block(lines):
    """Render a shaded code block."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, C_CODE_BG)
    cell.width = Inches(6.5)
    tf = cell.paragraphs[0]
    tf.paragraph_format.space_before = Pt(4)
    tf.paragraph_format.space_after  = Pt(4)
    first = True
    for line in lines:
        if first:
            p = tf
            first = False
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(line)
        run.font.name  = 'Courier New'
        run.font.size  = Pt(9.5)
        run.font.color.rgb = C_DARK
    doc.add_paragraph()  # spacing after

def note_box(text, kind='note'):
    colors = {
        'note':    (RGBColor(0xDB, 0xEA, 0xFE), '📘 Note'),
        'warning': (RGBColor(0xFF, 0xED, 0xD5), '⚠️  Warning'),
        'tip':     (RGBColor(0xD1, 0xFA, 0xE5), '✅ Tip'),
    }
    bg, label = colors.get(kind, colors['note'])
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    r1 = p.add_run(f'{label}:  ')
    r1.font.bold = True
    r1.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    doc.add_paragraph()

def spacer():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)

def simple_table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, C_HDR_BG)
        if col_widths:
            cell.width = Inches(col_widths[i])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.font.bold  = True
        run.font.size  = Pt(10)
        run.font.color.rgb = C_WHITE

    # Data rows
    for ri, row in enumerate(rows):
        tr = tbl.rows[ri + 1]
        bg = C_ROW_ALT if ri % 2 == 1 else C_WHITE
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            if col_widths:
                cell.width = Inches(col_widths[ci])
            p = cell.paragraphs[0]
            if isinstance(val, tuple):  # (text, bold)
                run = p.add_run(val[0])
                run.font.bold = val[1]
                run.font.size = Pt(10)
                run.font.color.rgb = C_DARK
            else:
                run = p.add_run(str(val))
                run.font.size  = Pt(10)
                run.font.color.rgb = C_DARK

    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run('MDM SaaS')
run.font.size  = Pt(36)
run.font.bold  = True
run.font.color.rgb = C_BLUE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Local Deployment Guide')
run2.font.size  = Pt(22)
run2.font.color.rgb = C_DARK

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(6)
run3 = p3.add_run('Step-by-Step Setup for macOS Development Environment')
run3.font.size  = Pt(13)
run3.font.color.rgb = C_GRAY

spacer()
spacer()

# Meta table
tbl = doc.add_table(rows=4, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'
meta = [
    ('Document Type', 'Deployment Guide'),
    ('Environment',   'Local Development (macOS)'),
    ('Version',       '1.0'),
    ('Last Updated',  'April 2026'),
]
for i, (k, v) in enumerate(meta):
    kc = tbl.rows[i].cells[0]
    vc = tbl.rows[i].cells[1]
    set_cell_bg(kc, C_CODE_BG)
    kc.width = Inches(2.2)
    vc.width = Inches(3.5)
    rk = kc.paragraphs[0].add_run(k)
    rk.font.bold = True
    rk.font.size = Pt(10)
    rk.font.color.rgb = C_DARK
    rv = vc.paragraphs[0].add_run(v)
    rv.font.size = Pt(10)
    rv.font.color.rgb = C_DARK

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════
h1('Table of Contents')
toc = [
    ('1.', 'Prerequisites',                              '3'),
    ('2.', 'Repository Setup',                           '3'),
    ('3.', 'Environment Configuration',                  '4'),
    ('4.', 'Generate Dev Certificates',                  '5'),
    ('5.', 'Start Docker Services',                      '6'),
    ('6.', 'Database Setup',                             '7'),
    ('7.', 'Start the Dashboard (Next.js)',               '7'),
    ('8.', 'Configure Public URL (ngrok)',                '8'),
    ('9.', 'Enroll Your First Device',                   '9'),
    ('10.', 'Install the Management Agent',              '10'),
    ('11.', 'Verify End-to-End',                         '11'),
    ('12.', 'Environment Variables Reference',           '12'),
    ('13.', 'Common Services & Ports',                   '13'),
    ('14.', 'Daily Dev Workflow',                        '14'),
]
for num, title, page in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f'{num}  {title}')
    r1.font.size = Pt(11)
    r1.font.color.rgb = C_DARK
    r2 = p.add_run(f'{"." * (55 - len(num) - len(title))}{page}')
    r2.font.size = Pt(11)
    r2.font.color.rgb = C_GRAY

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — PREREQUISITES
# ══════════════════════════════════════════════════════════════════════════════
h1('1.  Prerequisites')
body('Ensure the following are installed on your Mac before proceeding.')
spacer()

simple_table(
    ['Tool', 'Version', 'Install Command', 'Purpose'],
    [
        ('Docker Desktop', '4.x+',     'https://docker.com/products/docker-desktop', 'Runs app, DB, LocalStack'),
        ('Python',         '3.12+',    'brew install python@3.12',                   'Local scripts'),
        ('Node.js',        '18+',      'brew install node',                          'Next.js dashboard'),
        ('ngrok',          'Latest',   'brew install ngrok/ngrok/ngrok',             'Public HTTPS tunnel'),
        ('Git',            'Any',      'xcode-select --install',                     'Clone repository'),
        ('OpenSSL',        'Any',      'brew install openssl',                       'Generate certificates'),
    ],
    col_widths=[1.2, 0.8, 2.5, 1.9]
)

note_box(
    'Docker Desktop must be running before you start any docker compose commands. '
    'Allocate at least 4 GB RAM to Docker in Docker Desktop → Settings → Resources.',
    'warning'
)

h2('1.1  Verify Prerequisites')
code_block([
    '# Check all tools are available:',
    'docker --version          # Docker version 24.x or higher',
    'docker compose version    # Docker Compose v2.x',
    'python3 --version         # Python 3.12.x',
    'node --version            # v18.x or higher',
    'npm --version             # 9.x or higher',
    'ngrok version             # ngrok version 3.x',
    'openssl version           # OpenSSL 3.x',
])


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — REPOSITORY SETUP
# ══════════════════════════════════════════════════════════════════════════════
h1('2.  Repository Setup')

h2('2.1  Clone the Repository')
code_block([
    'git clone <repository-url> mdm-saas',
    'cd mdm-saas',
])

h2('2.2  Directory Structure Overview')
body('Key directories you will work with:')
simple_table(
    ['Path', 'Description'],
    [
        ('app/',                    'FastAPI backend — all Python source code'),
        ('app/api/routes/',         'API endpoint handlers'),
        ('app/mdm/apple/',          'Apple MDM protocol implementation'),
        ('app/db/',                 'SQLAlchemy models + Alembic migrations'),
        ('dashboard/',              'Next.js admin dashboard'),
        ('scripts/',                'Setup scripts, seed data, utilities'),
        ('certs/',                  'TLS and signing certificates (git-ignored)'),
        ('infra/terraform/',        'AWS infrastructure (ECS, RDS, SQS)'),
        ('.env',                    'Local environment variables (git-ignored)'),
        ('docker-compose.yml',      'Local dev service definitions'),
    ],
    col_widths=[2.2, 4.3]
)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — ENVIRONMENT CONFIGURATION
# ══════════════════════════════════════════════════════════════════════════════
h1('3.  Environment Configuration')

h2('3.1  Create the .env File')
code_block([
    'cp .env.example .env',
    'open -e .env   # or use any text editor',
])

h2('3.2  Minimum Required Variables')
body('Edit .env and set these values before starting:')
spacer()
simple_table(
    ['Variable', 'Value to Set', 'Notes'],
    [
        ('DATABASE_URL',         'postgresql+asyncpg://mdm:mdm@db:5432/mdmdb',  'Matches docker-compose DB service'),
        ('SECRET_KEY',           '<random 64-char hex string>',                  'Run: openssl rand -hex 32'),
        ('MDM_SERVER_URL',       'https://<your-ngrok-url>.ngrok-free.app',      'Set after starting ngrok (Step 8)'),
        ('APNS_CERT_PATH',       './certs/dev/apns.pem',                         'APNs push certificate'),
        ('APNS_KEY_PATH',        './certs/dev/apns.key',                         'APNs push private key'),
        ('APNS_USE_SANDBOX',     'false',                                         'false for production APNs cert'),
        ('APNS_PUSH_TOPIC',      'com.apple.mgmt.External.<uid>',                'From APNs cert UID field'),
        ('MDM_SIGNING_CERT_PATH','./certs/dev/mdm_signing.pem',                  'Profile signing cert'),
        ('MDM_SIGNING_KEY_PATH', './certs/dev/mdm_signing.key',                  'Profile signing key'),
        ('AWS_REGION',           'ap-south-1',                                   'Any region — LocalStack ignores it'),
        ('SQS_COMMAND_QUEUE_URL','http://localstack:4566/.../mdm-commands',       'Set after Step 5'),
    ],
    col_widths=[2.2, 2.3, 2.0]
)

h2('3.3  Generate a SECRET_KEY')
code_block([
    '# Generate a secure random key:',
    'openssl rand -hex 32',
    '# Copy the output and paste as SECRET_KEY in .env',
])


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — CERTIFICATES
# ══════════════════════════════════════════════════════════════════════════════
h1('4.  Generate Dev Certificates')

body('Two types of certificates are required:')
bullet('Device Identity Certificate  — embedded in enrollment profiles (PKCS#12 format)', bold_prefix='')
bullet('APNs Push Certificate  — authenticates the server with Apple Push Notification Service', bold_prefix='')
spacer()

h2('4.1  Device Identity Certificate (Required)')
body('This is a self-signed certificate used in the MDM enrollment profile. Run the generator script:')
code_block([
    '# Create certs directory:',
    'mkdir -p certs/dev',
    '',
    '# Generate device identity P12:',
    'bash scripts/gen_dev_certs.sh',
    '',
    '# Verify files were created:',
    'ls -la certs/dev/',
    '# Expected: device_identity.p12, device_identity.pem, device_identity.key',
])

h2('4.2  APNs Push Certificate (Required for Push Notifications)')
body('A real APNs MDM push certificate is required to wake devices. Follow these steps:')

numbered('Register at https://mdmcert.download using an organizational email address.')
numbered('Generate an encryption key pair on your machine:')
code_block([
    'openssl req -newkey rsa:2048 -keyout ./certs/dev/encrypt.key \\',
    '  -out ./certs/dev/encrypt.csr -subj "/CN=MDM Encrypt" -nodes',
])
numbered('Generate the APNs push certificate signing request (CSR):')
code_block([
    'openssl req -newkey rsa:2048 -keyout ./certs/dev/apns.key \\',
    '  -out ./certs/dev/apns.csr -subj "/CN=MDM Push" -nodes',
])
numbered('Submit the CSR to mdmcert.download and receive a signed .p7 response file.')
numbered('Decrypt the response using your encryption key:')
code_block([
    'openssl smime -decrypt -in response.p7 \\',
    '  -inkey ./certs/dev/encrypt.key -out push_request_decrypted.plist',
])
numbered('Upload the decrypted file to https://identity.apple.com/pushcert')
numbered('Download the certificate and save as ./certs/dev/apns.pem')
numbered('Extract the push topic from the certificate:')
code_block([
    'openssl x509 -in ./certs/dev/apns.pem -noout -subject | grep UID',
    '# Example output: UID=com.apple.mgmt.External.abc123...',
    '# Copy the full UID value into APNS_PUSH_TOPIC in .env',
])

h2('4.3  MDM Profile Signing Certificate (Optional for Dev)')
body('For development, you can skip profile signing. For signed profiles (required for supervised devices), '
     'use an Apple Developer ID Application certificate:')
code_block([
    '# Export from Keychain Access as .p12, then convert:',
    'openssl pkcs12 -in developer_id.p12 -nokeys -out ./certs/dev/mdm_signing.pem',
    'openssl pkcs12 -in developer_id.p12 -nocerts -nodes -out ./certs/dev/mdm_signing.key',
])
note_box('If MDM_SIGNING_CERT_PATH is not set, profiles are delivered unsigned. '
         'Unsigned profiles work on unsupervised devices for development.', 'note')


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — DOCKER SERVICES
# ══════════════════════════════════════════════════════════════════════════════
h1('5.  Start Docker Services')

h2('5.1  Start All Services')
code_block([
    '# Start all services (app, PostgreSQL, LocalStack):',
    'docker compose up -d',
    '',
    '# Verify all containers are running:',
    'docker compose ps',
    '',
    '# Expected output:',
    '# NAME                    STATUS',
    '# mdm-saas-app-1          Up (healthy)',
    '# mdm-saas-db-1           Up (healthy)',
    '# mdm-saas-localstack-1   Up (healthy)',
])

h2('5.2  Create the SQS Queue in LocalStack')
code_block([
    '# Create the MDM command queue in LocalStack:',
    'aws --endpoint-url=http://localhost:4566 sqs create-queue \\',
    '  --queue-name mdm-commands --region ap-south-1',
    '',
    '# Get the queue URL (copy this into .env as SQS_COMMAND_QUEUE_URL):',
    'aws --endpoint-url=http://localhost:4566 sqs get-queue-url \\',
    '  --queue-name mdm-commands --region ap-south-1',
    '',
    '# The URL will look like:',
    '# http://localhost:4566/000000000000/mdm-commands',
    '',
    '# Update SQS_COMMAND_QUEUE_URL in .env, then restart app:',
    'docker compose restart app',
])

h2('5.3  View Live Logs')
code_block([
    '# Watch app logs in real time:',
    'docker compose logs app -f',
    '',
    '# Watch all service logs:',
    'docker compose logs -f',
    '',
    '# Filter for errors only:',
    'docker compose logs app | grep -i "error\\|exception\\|traceback"',
])

note_box('The app container automatically restarts if it crashes. '
         'If it keeps restarting, check logs for missing environment variables.', 'warning')


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — DATABASE SETUP
# ══════════════════════════════════════════════════════════════════════════════
h1('6.  Database Setup')

h2('6.1  Run Migrations')
code_block([
    '# Apply all database migrations:',
    'docker compose exec app alembic upgrade head',
    '',
    '# Verify migration applied:',
    'docker compose exec app alembic current',
    '# Should show: <revision_id> (head)',
])

h2('6.2  Seed Initial Data')
code_block([
    '# Create the default tenant and admin user:',
    'docker compose exec app python scripts/seed_db.py',
    '',
    '# Default credentials created:',
    '# Email:    admin@acme.com',
    '# Password: secret',
    '# Tenant:   Acme Corp',
])

h2('6.3  Verify Database')
code_block([
    '# Connect to the database:',
    'docker compose exec -T db psql -U mdm mdmdb',
    '',
    '# Check tables exist:',
    '\\dt',
    '',
    '# Check tenant was created:',
    'SELECT name, slug, plan FROM tenants;',
    '',
    '# Check admin user:',
    'SELECT email, role FROM users;',
    '',
    '# Exit psql:',
    '\\q',
])


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — DASHBOARD
# ══════════════════════════════════════════════════════════════════════════════
h1('7.  Start the Dashboard (Next.js)')

h2('7.1  Install Dependencies')
code_block([
    'cd dashboard',
    'npm install',
])

h2('7.2  Start the Dev Server')
code_block([
    '# Still inside the dashboard/ directory:',
    'npm run dev',
    '',
    '# Dashboard will be available at:',
    '# http://localhost:3000',
])

h2('7.3  Configure API URL in Dashboard')
body('Open the dashboard in your browser at http://localhost:3000 and:')
numbered('Click the Settings icon (gear) in the sidebar.')
numbered('Set the API URL to your server URL (e.g. http://localhost:8000 for local, or your ngrok URL).')
numbered('Click Save.')
numbered('Log in with:  admin@acme.com  /  secret')
spacer()
note_box('The API URL is stored in browser localStorage. '
         'If you change your ngrok URL, update it in the dashboard Settings.', 'note')

h2('7.4  Verify Dashboard is Working')
body('After login you should see:')
bullet('Devices list (empty until enrollment)')
bullet('Sidebar with: Devices, Policies, Profiles, Packages, Compliance, Audit')
bullet('Settings page showing current API URL and tenant info')


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — NGROK
# ══════════════════════════════════════════════════════════════════════════════
h1('8.  Configure Public URL (ngrok)')

body('Apple MDM and APNs require a publicly accessible HTTPS URL. '
     'ngrok creates a secure tunnel from a public URL to your local server.')
spacer()

h2('8.1  Start ngrok')
code_block([
    '# Open a new terminal tab and run:',
    'ngrok http 8000',
    '',
    '# ngrok will display something like:',
    '# Forwarding  https://corking-ortho-joye.ngrok-free.app -> localhost:8000',
    '',
    '# Copy the https:// URL — this is your MDM_SERVER_URL',
])

h2('8.2  Update Configuration')
code_block([
    '# Edit .env and set:',
    'MDM_SERVER_URL=https://corking-ortho-joye.ngrok-free.app',
    '',
    '# Restart the app to pick up the new URL:',
    'docker compose restart app',
])

note_box(
    'The free ngrok URL changes every time you restart ngrok. '
    'When the URL changes: update MDM_SERVER_URL in .env, restart the app, '
    'and generate new enrollment tokens (old tokens point to the old URL). '
    'Consider ngrok paid plan for a fixed domain.',
    'warning'
)

h2('8.3  Verify the Public URL Works')
code_block([
    '# Test the API is reachable via ngrok:',
    'curl -s https://<your-ngrok-url>/api/v1/healthz',
    '# Expected: {"status": "ok"}',
    '',
    '# Test login:',
    'curl -s -X POST https://<your-ngrok-url>/api/v1/auth/login \\',
    '  -H "Content-Type: application/json" \\',
    '  -d \'{"email":"admin@acme.com","password":"secret"}\'',
    '# Expected: {"access_token": "eyJ..."}',
])

h2('8.4  Update Dashboard API URL')
body('After updating ngrok URL, also update the API URL in the dashboard:')
numbered('Open http://localhost:3000')
numbered('Go to Settings')
numbered('Update API URL to new ngrok URL')
numbered('Click Save and verify devices page loads without errors')


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — ENROLL DEVICE
# ══════════════════════════════════════════════════════════════════════════════
h1('9.  Enroll Your First Device')

h2('9.1  Generate an Enrollment Token')
body('In the dashboard:')
numbered('Go to Enrollment in the sidebar.')
numbered('Click Generate Token.')
numbered('Select platform: macOS.')
numbered('Choose Reusable if enrolling multiple devices.')
numbered('Click Create — an enrollment URL is shown.')
spacer()
body('Or via API:')
code_block([
    'curl -s -X POST https://<ngrok-url>/api/v1/enrollment/tokens \\',
    '  -H "Authorization: Bearer <your-jwt>" \\',
    '  -H "Content-Type: application/json" \\',
    '  -d \'{"platform":"macos","reusable":true,"expires_in_hours":72}\'',
])

h2('9.2  Open Enrollment URL on the Mac')
note_box('The enrollment URL MUST be opened in Safari. Chrome and Firefox cannot install MDM profiles.', 'warning')
numbered('Copy the enrollment URL from the dashboard.')
numbered('On the Mac to be enrolled, open Safari.')
numbered('Navigate to the enrollment URL.')
numbered('Safari prompts: "This website is trying to open System Preferences." — click Allow.')
numbered('System Settings → Privacy & Security → Profiles → downloaded profile appears.')
numbered('Click Install and enter the Mac\'s admin password when prompted.')
numbered('Click Install again to confirm.')

h2('9.3  Verify Enrollment')
code_block([
    '# Check macOS MDM status:',
    'sudo profiles status -type enrollment',
    '# Should show: MDM enrollment: Yes (User Approved)',
    '',
    '# List installed profiles:',
    'sudo profiles list',
])
body('In the dashboard, the device should appear in the Devices list within ~30 seconds '
     'with status enrolled.')
spacer()
note_box('If the device does not appear, check the app logs: '
         'docker compose logs app | grep -i "checkin\\|authenticate"', 'tip')


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — MANAGEMENT AGENT
# ══════════════════════════════════════════════════════════════════════════════
h1('10.  Install the Management Agent')

body('The management agent is a pure-bash LaunchDaemon that polls the server for jobs '
     '(software installs, admin access, scripts). It requires no Python or Xcode.')
spacer()

h2('10.1  Get the Bootstrap Command')
numbered('In the dashboard, open the device detail page.')
numbered('Scroll to the Agent Bootstrap section.')
numbered('Copy the bootstrap command — it looks like:')
code_block([
    'curl -sSLG -d auth=<agent_token> https://<ngrok-url>/api/v1/agent/bootstrap | sudo bash',
])

h2('10.2  Run the Bootstrap Command on the Mac')
note_box(
    'Important: Do NOT use double or single quotes around the URL or token when typing in Terminal. '
    'macOS Terminal auto-converts typed quotes to smart quotes which will break the command. '
    'Copy and paste the command directly from the dashboard.',
    'warning'
)
code_block([
    '# Paste exactly as shown in dashboard — no modification needed:',
    'curl -sSLG -d auth=<token> <bootstrap_url> | sudo bash',
    '',
    '# Enter your Mac admin password when prompted for sudo.',
])

h2('10.3  Verify Agent is Running')
code_block([
    '# Check LaunchDaemon is loaded:',
    'sudo launchctl list | grep mdmsaas',
    '# Expected: PID shown in first column',
    '',
    '# Watch agent logs:',
    'tail -f /var/log/mdm-agent.log',
    '# Expected: "MDM Agent started" and periodic "poll ok" messages',
])

h2('10.4  Test the Agent')
body('In the dashboard:')
numbered('Go to the device detail page.')
numbered('Click Query Device Info — the device should update within 30 seconds.')
numbered('Open the Self-Service Portal URL (shown in Agent section) in a browser on the Mac.')
numbered('The software catalog should appear.')


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 11 — VERIFY END-TO-END
# ══════════════════════════════════════════════════════════════════════════════
h1('11.  Verify End-to-End')

body('Run through this checklist to confirm everything is working:')
spacer()

simple_table(
    ['#', 'Test', 'Expected Result', 'Status'],
    [
        ('1',  'Dashboard login',               'JWT token issued, devices page loads',        '☐'),
        ('2',  'Device enrolled',               'Device shows in Devices list as "enrolled"',  '☐'),
        ('3',  'Query Device Info',             'Serial, OS, model updated in dashboard',      '☐'),
        ('4',  'APNs push received',            'Command executes within 30s of queuing',      '☐'),
        ('5',  'Push PSSO profile',             'Profile installs on Mac (sudo profiles list)','☐'),
        ('6',  'Push USB Block',                'USB Block profile on Mac, USB drive blocked', '☐'),
        ('7',  'Remove USB Block',              'Profile removed, USB drive mounts again',     '☐'),
        ('8',  'Agent polling',                 '/var/log/mdm-agent.log shows poll ok',        '☐'),
        ('9',  'Self-service portal',           'Catalog loads in browser on Mac',             '☐'),
        ('10', 'Software request + install',    'App installs after admin approval',           '☐'),
        ('11', 'JIT admin request',             'User elevated, auto-revoked after timeout',   '☐'),
        ('12', 'Compliance evaluation',         'Policy result shows in compliance dashboard', '☐'),
        ('13', 'Audit log',                     'All above actions appear in audit trail',     '☐'),
    ],
    col_widths=[0.3, 2.1, 2.6, 0.9]
)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 12 — ENV VARS REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
h1('12.  Environment Variables Reference')

simple_table(
    ['Variable', 'Required', 'Description', 'Example Value'],
    [
        ('DATABASE_URL',              'Yes', 'PostgreSQL async connection string',      'postgresql+asyncpg://mdm:mdm@db:5432/mdmdb'),
        ('SECRET_KEY',                'Yes', '256-bit JWT signing key (hex)',           'openssl rand -hex 32'),
        ('MDM_SERVER_URL',            'Yes', 'Public HTTPS URL (ngrok or fixed domain)','https://xyz.ngrok-free.app'),
        ('APNS_CERT_PATH',            'Yes', 'Path to APNs push certificate PEM',      './certs/dev/apns.pem'),
        ('APNS_KEY_PATH',             'Yes', 'Path to APNs push private key',          './certs/dev/apns.key'),
        ('APNS_USE_SANDBOX',          'Yes', 'Use APNs sandbox vs production',         'false'),
        ('APNS_PUSH_TOPIC',           'Yes', 'APNs push topic from certificate UID',   'com.apple.mgmt.External.xxx'),
        ('DEVICE_IDENTITY_P12_PATH',  'Yes', 'PKCS#12 device identity cert path',      './certs/dev/device_identity.p12'),
        ('MDM_SIGNING_CERT_PATH',     'No',  'Profile signing certificate path',       './certs/dev/mdm_signing.pem'),
        ('MDM_SIGNING_KEY_PATH',      'No',  'Profile signing private key path',       './certs/dev/mdm_signing.key'),
        ('SQS_COMMAND_QUEUE_URL',     'Yes', 'SQS queue URL for MDM commands',         'http://localstack:4566/.../mdm-commands'),
        ('AWS_REGION',                'Yes', 'AWS region (LocalStack ignores this)',   'ap-south-1'),
        ('AWS_ACCESS_KEY_ID',         'Dev', 'Dummy value for LocalStack',             'test'),
        ('AWS_SECRET_ACCESS_KEY',     'Dev', 'Dummy value for LocalStack',             'test'),
        ('NOTIFICATION_WEBHOOK_URL',  'No',  'Slack/Teams/Discord webhook for alerts', 'https://hooks.slack.com/...'),
        ('UPLOAD_DIR',                'No',  'Package upload directory',               '/app/uploads/packages'),
        ('ENVIRONMENT',               'No',  'Runtime environment name',               'development'),
    ],
    col_widths=[2.1, 0.7, 2.1, 1.55]
)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 13 — SERVICES & PORTS
# ══════════════════════════════════════════════════════════════════════════════
h1('13.  Common Services & Ports')

simple_table(
    ['Service', 'Local URL', 'Description'],
    [
        ('MDM API (FastAPI)',        'http://localhost:8000',         'Main backend API server'),
        ('API Docs (Swagger)',       'http://localhost:8000/docs',    'Interactive API documentation'),
        ('API Docs (ReDoc)',         'http://localhost:8000/redoc',   'Alternative API documentation'),
        ('Next.js Dashboard',       'http://localhost:3000',         'Admin web dashboard'),
        ('PostgreSQL',              'localhost:5433',                'Database (port 5433 to avoid conflicts)'),
        ('LocalStack (SQS)',        'http://localhost:4566',         'AWS service emulator'),
        ('ngrok Dashboard',         'http://localhost:4040',         'ngrok tunnel inspector & request log'),
    ],
    col_widths=[2.0, 2.2, 2.3]
)

h2('13.1  Useful psql Commands')
code_block([
    '# Connect to database:',
    'docker compose exec -T db psql -U mdm mdmdb',
    '',
    '# List all tables:',
    '\\dt',
    '',
    '# Describe a table:',
    '\\d devices',
    '',
    '# Common queries:',
    'SELECT hostname, status, last_checkin FROM devices;',
    'SELECT command_type, status, queued_at FROM mdm_commands ORDER BY queued_at DESC LIMIT 10;',
    'SELECT actor_email, action, created_at FROM audit_logs ORDER BY created_at DESC LIMIT 10;',
])


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 14 — DAILY DEV WORKFLOW
# ══════════════════════════════════════════════════════════════════════════════
h1('14.  Daily Dev Workflow')

h2('14.1  Starting Up Each Day')
code_block([
    '# 1. Start Docker services:',
    'docker compose up -d',
    '',
    '# 2. Start ngrok (new terminal tab):',
    'ngrok http 8000',
    '',
    '# 3. Update MDM_SERVER_URL in .env with new ngrok URL',
    '#    (skip if using paid fixed ngrok domain)',
    '',
    '# 4. Restart app to pick up new URL:',
    'docker compose restart app',
    '',
    '# 5. Update API URL in dashboard Settings',
    '#    (if ngrok URL changed)',
    '',
    '# 6. Start dashboard (new terminal tab):',
    'cd dashboard && npm run dev',
])

h2('14.2  After Changing Python Code')
code_block([
    '# Restart the app container:',
    'docker compose restart app',
    '',
    '# Watch logs to confirm no errors:',
    'docker compose logs app -f',
])

h2('14.3  After Changing DB Models')
code_block([
    '# Generate a new migration:',
    'docker compose exec app alembic revision --autogenerate -m "your description"',
    '',
    '# Review the generated file in app/db/migrations/versions/',
    '',
    '# Apply the migration:',
    'docker compose exec app alembic upgrade head',
])

h2('14.4  Shutting Down')
code_block([
    '# Stop all Docker services (data is preserved):',
    'docker compose down',
    '',
    '# Stop and delete all data (fresh start):',
    'docker compose down -v',
    '',
    '# Kill ngrok: Ctrl+C in the ngrok terminal',
    '# Kill dashboard: Ctrl+C in the npm run dev terminal',
])

h2('14.5  Force a Device to Check In')
code_block([
    '# On the enrolled Mac — triggers immediate MDM check-in:',
    'sudo /usr/libexec/mdmclient daemon',
    '',
    '# Useful when APNs push is delayed (common on VMs)',
])

h2('14.6  Reset Everything for a Clean Start')
code_block([
    '# Stop containers and delete all data:',
    'docker compose down -v',
    '',
    '# Start fresh:',
    'docker compose up -d',
    'docker compose exec app alembic upgrade head',
    'docker compose exec app python scripts/seed_db.py',
    '',
    '# Re-create SQS queue:',
    'aws --endpoint-url=http://localhost:4566 sqs create-queue \\',
    '  --queue-name mdm-commands --region ap-south-1',
    '',
    '# On enrolled Mac — remove old MDM profile and re-enroll:',
    '# System Settings → Privacy & Security → Profiles → Remove MDM profile',
    '# Then generate new enrollment token and re-enroll',
])

spacer()
spacer()
note_box(
    'For issues not covered here, refer to TROUBLESHOOTING.md in the repository root. '
    'It covers every known error with step-by-step fixes.',
    'tip'
)


# ══════════════════════════════════════════════════════════════════════════════
out = '/Users/manikandank/Downloads/mdm-saas/MDM_SaaS_Local_Deployment_Guide.docx'
doc.save(out)
print(f'Saved: {out}')
